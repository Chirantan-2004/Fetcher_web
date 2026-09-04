from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from app.schemas import CanonicalDocument

def generate_docx(document: CanonicalDocument) -> bytes:
    output = BytesIO(); doc = Document(); section = doc.sections[0]; section.top_margin = Inches(.7); section.bottom_margin = Inches(.7)
    doc.add_heading(document.title or "Fetched document", 0); doc.add_paragraph(document.source_url, style="Subtitle")
    doc.add_heading("Metadata", 1); doc.add_paragraph(f"Domain: {document.domain}\nAuthor: {document.author or 'Unknown'}\nWords: {document.statistics.word_count}\nFetched: {document.fetched_at.isoformat()}")
    doc.add_heading("Extracted Content", 1)
    for section_data in document.content["sections"]:
        if section_data.heading: doc.add_heading(section_data.heading, min(section_data.level, 9))
        for block in section_data.blocks:
            if block.type == "paragraph": doc.add_paragraph(block.text or "")
            elif block.type == "quote": doc.add_paragraph(block.text or "", style="Intense Quote")
            elif block.type == "code": doc.add_paragraph(block.text or "", style="No Spacing")
            elif block.type == "list":
                for item in block.items: doc.add_paragraph(item, style="List Bullet")
            elif block.type == "table":
                table = doc.add_table(rows=1, cols=max(1, len(block.headers))); table.style = "Light Shading Accent 1"
                for cell, value in zip(table.rows[0].cells, block.headers): cell.text = value
                for row in block.rows:
                    cells = table.add_row().cells
                    for cell, value in zip(cells, row): cell.text = value
    doc.add_heading("Source Information", 1); doc.add_paragraph(f"Original URL: {document.source_url}\nCanonical URL: {document.canonical_url}\nFetcher schema: {document.schema_version}")
    doc.save(output); return output.getvalue()
